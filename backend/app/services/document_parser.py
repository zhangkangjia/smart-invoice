"""文档解析服务：从 Word/PDF 中提取文本。

提取后走文字识别（text_recognizer）链路。
"""

import logging
from io import BytesIO
from typing import Optional

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = (".docx", ".doc", ".pdf", ".txt")


def extract_text_from_document(file_bytes: bytes, filename: str) -> str:
    """从文档中提取纯文本。

    支持 .docx / .pdf / .txt
    """
    lower = filename.lower()
    try:
        if lower.endswith(".docx"):
            return _extract_from_docx(file_bytes)
        if lower.endswith(".pdf"):
            return _extract_from_pdf(file_bytes)
        if lower.endswith(".txt"):
            return file_bytes.decode("utf-8", errors="ignore")
        if lower.endswith(".doc"):
            # 老 .doc 格式不支持（需要 libreoffice 转换），提示用户
            raise ValueError("不支持 .doc 老格式，请另存为 .docx 后再上传")
        raise ValueError(f"不支持的文件类型: {filename}")
    except ValueError:
        raise
    except Exception as e:
        logger.exception("文档解析失败: %s", filename)
        raise ValueError(f"文档解析失败: {e}")


def _extract_from_docx(file_bytes: bytes) -> str:
    """从 .docx 提取文本（段落 + 表格）。"""
    from docx import Document

    doc = Document(BytesIO(file_bytes))
    parts: list[str] = []

    # 1. 普通段落
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # 2. 表格内容
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)


def _extract_from_pdf(file_bytes: bytes) -> str:
    """从 PDF 提取文本（所有页）。"""
    import pdfplumber

    parts: list[str] = []
    with pdfplumber.open(BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            if text.strip():
                parts.append(text.strip())

            # 提取表格
            for table in page.extract_tables() or []:
                for row in table:
                    cells = [str(cell).strip() for cell in row if cell]
                    if cells:
                        parts.append(" | ".join(cells))

    return "\n".join(parts)
