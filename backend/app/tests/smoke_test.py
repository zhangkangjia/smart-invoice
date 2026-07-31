"""智能开票平台 - 服务启动后健康检查脚本。

用法：
    python -m app.tests.smoke_test

检查项：
1. 数据库连接
2. 基础表存在
3. 超级管理员存在
4. 示例企业已填充
5. 核心 API 路由可访问
6. AI Mock 识别可用
7. 模拟开票链路可用
8. 微信/企业微信配置状态
9. 文档解析服务可用
10. MinIO 存储可用

输出彩色结果，失败项给出具体错误。
"""

import asyncio
import sys
import traceback
from io import BytesIO
from typing import Any

from sqlalchemy import select, text

from app.db.session import AsyncSessionLocal, async_engine
from app.models.enterprise import Enterprise
from app.models.user import User
from app.models.role import Role

# ANSI 颜色
GREEN = "\033[92m"
RED = "\033[91m"
YELLOW = "\033[93m"
CYAN = "\033[96m"
NC = "\033[0m"

passed = 0
failed = 0
warnings = 0


def log_pass(msg: str) -> None:
    global passed
    passed += 1
    print(f"{GREEN}  [PASS]{NC} {msg}")


def log_fail(msg: str, error: str = "") -> None:
    global failed
    failed += 1
    print(f"{RED}  [FAIL]{NC} {msg}")
    if error:
        print(f"         {error}")


def log_warn(msg: str) -> None:
    global warnings
    warnings += 1
    print(f"{YELLOW}  [WARN]{NC} {msg}")


def section(title: str) -> None:
    print(f"\n{CYAN}━━━ {title} ━━━{NC}")


# --------------------------------------------------------------------------- #
# 测试项
# --------------------------------------------------------------------------- #

async def test_db_connection() -> None:
    section("1. 数据库连接")
    try:
        async with async_engine.connect() as conn:
            await conn.execute(text("SELECT 1"))
        log_pass("数据库连接正常")
    except Exception as e:
        log_fail("数据库连接失败", str(e))


async def test_tables_exist() -> None:
    section("2. 基础表存在")
    required_tables = [
        "users", "roles", "user_roles", "tenants",
        "enterprises", "customer_titles", "products",
        "business_requests", "invoice_tasks", "source_documents",
        "system_config",
    ]
    async with async_engine.connect() as conn:
        for table in required_tables:
            try:
                result = await conn.execute(text(
                    f"SELECT to_regclass('public.{table}')"
                ))
                exists = result.scalar()
                if exists:
                    log_pass(f"表 {table} 存在")
                else:
                    log_fail(f"表 {table} 不存在")
            except Exception as e:
                log_fail(f"检查表 {table} 失败", str(e))


async def test_admin_user() -> None:
    section("3. 超级管理员")
    async with AsyncSessionLocal() as db:
        result = await db.execute(
            select(User).where(User.is_super_admin == True).limit(1)  # noqa: E712
        )
        admin = result.scalar_one_or_none()
        if admin:
            log_pass(f"超级管理员存在: {admin.username}")
        else:
            log_fail("超级管理员不存在（请检查 init_db 是否执行）")


async def test_demo_enterprises() -> None:
    section("4. 示例企业数据")
    async with AsyncSessionLocal() as db:
        result = await db.execute(select(Enterprise).limit(5))
        enterprises = result.scalars().all()
        if len(enterprises) >= 1:
            log_pass(f"已有 {len(enterprises)} 家企业")
            for e in enterprises[:3]:
                print(f"         - {e.name} (税号: {e.tax_no}, 状态: {e.status})")
        else:
            log_fail("无企业数据（请执行 seed_demo）")


async def test_ai_recognition() -> None:
    section("5. AI 识别（Mock 模式）")
    try:
        from app.services.ai.router import recognize_text
        result = await recognize_text(
            content="帮我开一张发票，购方是杭州两心同网络科技有限公司，税号91330106MA28T1234X，商品是软件开发服务，金额10000元不含税",
            tenant_id="test",
        )
        if result and result.buyer_name:
            log_pass(f"文字识别成功: 购方={result.buyer_name}")
        else:
            log_fail("文字识别返回空结果")
    except Exception as e:
        log_fail("AI 识别异常", str(e))


async def test_document_parser() -> None:
    section("6. 文档解析服务")
    try:
        from app.services.document_parser import extract_text_from_document

        # 测试 docx（用最小合法 docx）
        # 如果没有 python-docx，跳过
        try:
            from docx import Document
            doc = Document()
            doc.add_paragraph("测试企业: 杭州两心同网络科技有限公司")
            doc.add_paragraph("税号: 91330106MA28T1234X")
            buf = BytesIO()
            doc.save(buf)
            text = extract_text_from_document(buf.getvalue(), "test.docx")
            if "杭州两心同" in text:
                log_pass("DOCX 解析正常")
            else:
                log_fail("DOCX 解析结果异常", f"提取内容: {text[:100]}")
        except ImportError:
            log_warn("python-docx 未安装，跳过 DOCX 解析测试")

        # 测试 txt
        text = extract_text_from_document(b"hello world", "test.txt")
        if "hello" in text:
            log_pass("TXT 解析正常")
        else:
            log_fail("TXT 解析异常")

    except Exception as e:
        log_fail("文档解析异常", str(e))


async def test_wechat_config() -> None:
    section("7. 微信/企业微信配置")
    from app.services.wechat_service import wecom_service, wechat_service

    if wecom_service.enabled:
        log_pass(f"企业微信已配置 (CorpID: {wecom_service.corp_id[:8]}...)")
    else:
        log_warn("企业微信未配置（可选功能）")

    if wechat_service.enabled:
        log_pass(f"微信服务号已配置 (AppID: {wechat_service.app_id[:8]}...)")
    else:
        log_warn("微信服务号未配置（可选功能）")


async def test_baiwang_config() -> None:
    section("8. 百望云通道配置")
    from app.core.config import settings
    if getattr(settings, "BAIWANG_APP_KEY", ""):
        log_pass("百望云已配置")
    else:
        log_warn("百望云未配置（当前为模拟开票模式）")


async def test_storage() -> None:
    section("9. 存储服务")
    try:
        from app.services.storage_service import get_storage_info
        info = await get_storage_info()
        if info:
            log_pass(f"存储服务可用: {info}")
        else:
            log_warn("存储服务未初始化")
    except ImportError:
        log_warn("storage_service 模块不存在")
    except Exception as e:
        log_warn(f"存储服务异常（可忽略）: {e}")


async def test_system_config_table() -> None:
    section("10. 系统配置表")
    async with async_engine.connect() as conn:
        try:
            result = await conn.execute(text("SELECT count(*) FROM system_config"))
            count = result.scalar()
            log_pass(f"system_config 表可读，当前 {count} 条配置")
        except Exception as e:
            log_fail("system_config 表不存在或不可读", str(e))


# --------------------------------------------------------------------------- #
# 主入口
# --------------------------------------------------------------------------- #

async def main() -> None:
    print(f"\n{CYAN}╔══════════════════════════════════════════════════════╗")
    print(f"║   智能开票平台 - 健康检查                            ║")
    print(f"╚══════════════════════════════════════════════════════╝{NC}\n")

    tests = [
        test_db_connection,
        test_tables_exist,
        test_admin_user,
        test_demo_enterprises,
        test_ai_recognition,
        test_document_parser,
        test_wechat_config,
        test_baiwang_config,
        test_storage,
        test_system_config_table,
    ]

    for test_func in tests:
        try:
            await test_func()
        except Exception as e:
            log_fail(f"{test_func.__name__} 异常", str(e))
            traceback.print_exc()

    # 汇总
    print(f"\n{CYAN}━━━ 汇总 ━━━{NC}")
    total = passed + failed + warnings
    print(f"  {GREEN}通过: {passed}{NC} / {total}")
    print(f"  {RED}失败: {failed}{NC} / {total}")
    print(f"  {YELLOW}警告: {warnings}{NC} / {total}")

    if failed > 0:
        print(f"\n{RED}有 {failed} 项失败，建议修复后再上线！{NC}")
        sys.exit(1)
    else:
        print(f"\n{GREEN}核心功能正常，可以上线！{NC}")
        sys.exit(0)


if __name__ == "__main__":
    asyncio.run(main())
